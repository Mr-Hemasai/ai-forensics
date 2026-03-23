from __future__ import annotations

from pathlib import Path
import tempfile
from typing import Any

import matplotlib

matplotlib.use("Agg")

from app.core.store import CaseRecord
from app.services.analysis import analyze_frequency, build_case_profile, detect_suspicious_patterns, detect_time_patterns


def _prepare_case_datasets(case_record: CaseRecord) -> list[dict[str, Any]]:
    return [
        {"file_name": dataset.file_name, "dataframe": dataset.dataframe, "summary": dataset.summary}
        for dataset in case_record.datasets
    ]


def _create_frequency_chart(chart_data: list[dict[str, Any]], output_dir: Path) -> Path | None:
    if not chart_data:
        return None
    import matplotlib.pyplot as plt

    labels = [item["label"] for item in chart_data[:8]]
    values = [item["value"] for item in chart_data[:8]]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar(labels, values, color="#22d3ee")
    ax.set_title("Top Entities")
    ax.tick_params(axis="x", rotation=35)
    fig.tight_layout()
    output = output_dir / "frequency_chart.png"
    fig.savefig(output, dpi=150)
    plt.close(fig)
    return output


def _create_timeline_chart(timeline: list[dict[str, Any]], output_dir: Path) -> Path | None:
    if not timeline:
        return None
    import matplotlib.pyplot as plt

    buckets: dict[str, int] = {}
    for item in timeline:
        buckets[item["bucket"]] = buckets.get(item["bucket"], 0) + int(item["value"])
    labels = list(buckets.keys())
    values = list(buckets.values())
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(labels, values, color="#f59e0b", marker="o")
    ax.fill_between(range(len(values)), values, color="#f59e0b", alpha=0.2)
    ax.set_title("Activity Timeline")
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=35)
    fig.tight_layout()
    output = output_dir / "timeline_chart.png"
    fig.savefig(output, dpi=150)
    plt.close(fig)
    return output


def build_case_report(case_record: CaseRecord) -> Path:
    from docx import Document
    from docx.shared import Inches

    datasets = _prepare_case_datasets(case_record)
    case_profile = build_case_profile(datasets)
    frequency = analyze_frequency(datasets)
    suspicious = detect_suspicious_patterns(datasets)
    time_patterns = detect_time_patterns(datasets)

    temp_dir = Path(tempfile.mkdtemp(prefix="aiforensic-report-"))
    document = Document()
    document.add_heading(f"Investigation Report: {case_record.case_name}", level=0)
    document.add_paragraph(f"Case ID: {case_record.case_id}")
    document.add_paragraph(f"Datasets loaded: {len(case_record.datasets)}")

    document.add_heading("1. Case Details", level=1)
    document.add_paragraph(
        "This report summarizes uploaded telecom and digital forensic datasets, investigation queries, and detected leads."
    )

    document.add_heading("2. Dataset Summary", level=1)
    for dataset in case_record.datasets:
        summary = dataset.summary
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{summary['file_name']} ").bold = True
        paragraph.add_run(
            f"- {summary['dataset_type_guess']}, {summary['rows']} rows, {len(summary['columns'])} columns"
        )
        document.add_paragraph(
            f"Detected semantics: {', '.join(summary['semantic_summary']['semantic_types'].keys()) or 'none'}"
        )

    document.add_heading("3. Investigation Summary", level=1)
    document.add_paragraph(f"Available semantic types: {', '.join(case_profile['semantic_inventory'].keys()) or 'none'}")
    if frequency["top_entities"]:
        document.add_paragraph(
            f"Most active entity: {frequency['top_entities'][0]['value']} ({frequency['top_entities'][0]['count']} records)"
        )

    document.add_heading("4. Chat Analysis", level=1)
    for index, message in enumerate(case_record.chat_history, start=1):
        prefix = "Investigator" if message["role"] == "user" else "Assistant"
        document.add_paragraph(f"{prefix} {index}: {message['content']}")

    document.add_heading("5. Observation Box", level=1)
    for item in case_record.observation_items:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("6. Insights & Patterns", level=1)
    for lead in suspicious.get("leads", [])[:5]:
        document.add_paragraph(lead, style="List Bullet")
    for alert in suspicious.get("alerts", [])[:5]:
        document.add_paragraph(alert, style="List Bullet")

    document.add_heading("7. Leads & Alerts", level=1)
    if not suspicious.get("leads") and not suspicious.get("alerts"):
        document.add_paragraph("No strong lead or alert threshold was triggered in this case.")
    for lead in suspicious.get("leads", [])[:5]:
        document.add_paragraph(f"[LEAD] {lead}")
    for alert in suspicious.get("alerts", [])[:5]:
        document.add_paragraph(f"[ALERT] {alert}")

    document.add_heading("8. Visualizations", level=1)
    freq_image = _create_frequency_chart(frequency.get("visualizations", {}).get("frequency_chart", []), temp_dir)
    timeline_image = _create_timeline_chart(time_patterns.get("visualizations", {}).get("timeline", []), temp_dir)
    for image_path in [freq_image, timeline_image]:
        if image_path:
            document.add_picture(str(image_path), width=Inches(6.5))

    report_path = temp_dir / f"{case_record.case_name.replace(' ', '_').lower()}_report.docx"
    document.save(report_path)
    return report_path
